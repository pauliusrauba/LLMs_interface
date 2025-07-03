import pandas as pd
import pickle
import numpy as np
import ast
import matplotlib.pyplot as plt
import streamlit as st
import PIL
import io
from langchain.utilities import WikipediaAPIWrapper
from langchain.agents import tool
from autoprognosis.utils.serialization import load_from_file
from resources.model.QRisk_model import QRisk3Model
import shap

from langchain.vectorstores import FAISS
#from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.llms import AzureOpenAI
from langchain import PromptTemplate, LLMChain
from langchain.document_loaders import PyPDFLoader
from autoprognosis.plugins.explainers import Explainers
from openai_config import get_openai_config
import os
from docx import Document
from langchain.chat_models import AzureChatOpenAI

from langchain.embeddings import OpenAIEmbeddings
import dill
import re

config = get_openai_config()

import pickle
import io
import urllib
import base64

def person_data(name: str) -> pd.DataFrame:
    """Returns the data of a client. Use this when there is a need to 
    extract data of a patient. The function returns a pandas dataframe."""
    return pd.read_csv("./cvd/person_cvd.csv")

def bin_points(scores, bin_edges):
    assert(bin_edges is not None), "Bins have not been defined"
    scores = scores.squeeze()
    assert(np.size(scores.shape) < 2), "scores should be a 1D vector or singleton"
    scores = np.reshape(scores, (scores.size, 1))
    bin_edges = np.reshape(bin_edges, (1, bin_edges.size))
    return np.sum(scores > bin_edges, axis=1)

@tool
def get_feature_importance(name: str) -> str:
    """
    Use this for any question related to explaining what are the main factors 
    driving the risk of disease for any specific person without asking for any plot. If a plot is asked, to not use this.
    The input to this function is the name of the person (e.g. 'john') or the index 
    of a person (e.g. 'person_cvd_0', 'person_cvd_1', etc). 
    The function returns the top three factors contributing to the risk and their scores.
    It also generates a plot of feature importance.
    """

    # Load the necessary data and background dataset
    background = pd.read_csv("./resources/patient_info/background_dataset.csv").drop('target', axis=1)

    # Load the input data for the specified person
    # Assuming the person's data is saved as 'name.csv'
    X = pd.read_csv(f"./resources/patient_info/{name.lower()}.csv").T
    ap2_features = ['sex',  'age',  'b_atrial_fibr',  'b_steroid_treat',  'b_diab_type2',  'family_cvd',  'sbp',  'smallbin',  'alkaline_phosphatase',  'apolipoprotein_a',  'apolipoprotein_b',  'urea',  'c_reactive_protein',  'cystatin_c',  'glycated_haemoglobin',  'igf_0',  'lipoprotein_a',  'triglycerides',  'overall_health_rating',  'ht_treat']

    X.columns = ap2_features

    # Load ap2 as model
    with open('./resources/model/ap_model.pkl', 'rb') as f:
        model = dill.load(f)

    # SHAP analysis
    explainer = shap.KernelExplainer(
        model.predict_proba, 
        background,
        link="logit",
        n_samples=1000,
        feature_names=X.columns.tolist() 
    )
    shap_values = explainer.shap_values(X)

    # Print results
    feature_importance = pd.DataFrame({
        'feature': X.columns,
        'importance': np.abs(shap_values[0]),
        'actual_value': X.iloc[0],
        'mean_background': background.mean()
    })
    feature_importance = feature_importance.sort_values('importance', ascending=False)

    return str(feature_importance)

@tool
def plot_feature_importance_heart_risk(name: str) -> list:
    """Use this for any question related to plotting the feature importance of heart risk for any patient or any model.
    The input should always be an empty string and this function will always return a tuple that contains the top three risks
    and their associated scores. It will always plot of feature importances. """
    # Assume that clf is the trained random forest model and X_train is the training data
    # Import the data for the person

    with open('./resources/patient_info/john_shap.pickle', 'rb') as file:
        shap_values = dill.load(file)

    shap_values_instance = shap_values.values

    feature_names = np.array(shap_values.feature_names)
    # Get the indices of the top 5 absolute SHAP values
    top_indices = np.argsort(np.abs(shap_values_instance))[-5:]

    # Use these indices to get the corresponding feature names
    top_feature_names = feature_names[top_indices]

    # Since argsort sorts in ascending order, we might want the top features in descending order of their absolute impact
    top_feature_names = top_feature_names[::-1]

    return "Top 5 feature names based on SHAP values:", top_feature_names


@tool
def get_information_on_patient(feature: str) -> str:
    """Use this function to extract a specific piece of information on the patient that is available in the pandas
    dataframe. This function is only used for extracting a single piece of information, not describing the patient.
    The information on the patient that is available is: ['age', 'sex', 'cp', 'trestbps', 'chol', 'fbs',
      'restecg', 'thalch', 'exang', 'oldpeak', 'slope', 'ca', 'thal'].
    The input to the function is the column name, and the output is a string explaining the value of that
    column name"""
    
    df = person_data('')
    
    info_person = df.loc[0, feature]
    
    return f"The value of {feature} is {info_person}"

@tool
def get_info_from_wkipedia(item: str) -> str:
    """Use this tool for any questions related to overall medical literature and overall knowledge,
    as well as extracting relevant statistics for diseases. The input for this tool is the object of search,
    such as a disease, and the output is wikipedia information for that disease. The language model uses
    this information to answer questions relevant to the person."""
    
    wiki = WikipediaAPIWrapper()
    return wiki.run(item)

@tool
def counterfactual_CVD_risk(features: str) -> str:
    """Use this for any question related to how the cardiovascular risk would change if any of the observed
    characteristics, such as age, would change. Use this for concrete calculations for a patient. The current columns are: 
    'sex': 'Biological gender',
    'age': 'Individuals age',
    'b_atrial_fibr': 'Atrial fibrillation presence',
    'b_antipsychotic_use': 'Antipsychotic medication use',
    'b_steroid_treat': 'Steroid treatment',
    'b_erectile_disf': 'Erectile dysfunction presence',
    'b_had_migraine': 'History of migraines',
    'b_rheumatoid_arthritis': 'Rheumatoid arthritis presence',
    'b_renal': 'Renal issues presence',
    'b_mental_illness': 'Mental illness presence',
    'b_sle': 'Systemic lupus erythematosus presence',
    'hypdbin': 'Hypertension presence',
    'b_diab_type1': 'Type 1 diabetes presence',
    'hxdiab': 'Diabetes history',
    'bmi': 'Body mass index',
    'ethrisk': 'Ethnicity risk factor',
    'family_cvd': 'Family cardiovascular disease history',
    'chol_ratio': 'Cholesterol ratio',
    'sbp': 'Systolic blood pressure',
    'sbps5': 'SBP standard deviation over 5 years',
    'smallbin': 'Small body size indicator',
    'town_depr_index': 'Town deprivation index'

    The function changes the required characteristic to the set value and re-runs the risk prediction.
    The function takes a string in the form tuple as an input which is "(feature, value)", such as "('sbp', 130)". 
    The function then returns a string explaining the old and new risk predictions, as well as their difference."""
    
    # Get data
    X = pd.read_csv('./resources/patient_info/person_cvd.csv')  
    
    feat, value = ast.literal_eval(features)
    X_count = X.copy()
    X_count[feat] = value
    
    # Get classifier
    qrisk_model = QRisk3Model()

    score_old = qrisk_model.predict(X).values[0][0].round(3)
    score_new = qrisk_model.predict(X_count).values[0][0].round(3)

    diff = score_new - score_old

    return score_old, score_new, diff

@tool
def counterfactual_CVD_risk_ap(input_data: str) -> str:
    """
    Use this function to calculate how the cardiovascular risk would change if any of the observed
    characteristics, such as age, would change for a specific patient.

    **Input Structure:**
    The function accepts a single string `input_data` in the following format:
    ```
    "person_name: 'person_cvd_0', feature_change: ('sbp', 130)"
    ```
    - `person_name`: The identifier for the patient's data file.
    - `feature_change`: A tuple representing the feature to change and its new value.

    **Output Structure:**
    The function returns a string explaining:
    - The original risk prediction.
    - The new risk prediction after the feature change.
    - The difference between the two predictions.

    **Current Columns:**
    ```
    'sex': 'Biological gender',
    'age': "Individual's age",
    'b_atrial_fibr': 'Atrial fibrillation presence',
    'b_steroid_treat': 'Steroid treatment',
    'b_diab_type2': 'Type 2 diabetes presence',
    'family_cvd': 'Family cardiovascular disease history',
    'sbp': 'Systolic blood pressure',
    'smallbin': 'Smoking status indicator',
    'alkaline_phosphatase': 'Alkaline phosphatase levels',
    'apolipoprotein_a': 'Apolipoprotein A levels',
    'apolipoprotein_b': 'Apolipoprotein B levels',
    'urea': 'Urea levels',
    'c_reactive_protein': 'C-reactive protein levels',
    'cystatin_c': 'Cystatin C levels',
    'glycated_haemoglobin': 'Glycated haemoglobin levels',
    'igf_0': 'Insulin-like growth factor 0 levels',
    'lipoprotein_a': 'Lipoprotein A levels',
    'triglycerides': 'Triglycerides levels',
    'overall_health_rating': 'Overall health rating',
    'ht_treat': 'Hypertension treatment presence'
    ```

    **Example Usage:**
    ```
    result = counterfactual_CVD_risk("person_name: 'person_cvd_0', feature_change: ('sbp', 130)")
    ```
    """

    import pandas as pd
    import ast
    import dill
    ap2_features = ['sex',  'age',  'b_atrial_fibr',  'b_steroid_treat',  'b_diab_type2',  'family_cvd',  'sbp',  'smallbin',  'alkaline_phosphatase',  'apolipoprotein_a',  'apolipoprotein_b',  'urea',  'c_reactive_protein',  'cystatin_c',  'glycated_haemoglobin',  'igf_0',  'lipoprotein_a',  'triglycerides',  'overall_health_rating',  'ht_treat']

    # Parse the input_data string to extract person_name and feature_change
    match = re.match(r"person_name:\s*'(.*?)',\s*feature_change:\s*(\(.+\))", input_data)
    if not match:
        return "Invalid input format. Please provide input in the format: \"person_name: 'name', feature_change: (feature, value)\""

    person_name = match.group(1)
    feature_change_str = match.group(2)

    # Parse the feature_change tuple
    try:
        feat, value = ast.literal_eval(feature_change_str)
    except (SyntaxError, ValueError):
        return "Invalid feature_change format. It should be a tuple, e.g., ('sbp', 130)"

    # Load the patient's data
    try:
        X = pd.read_csv(f'./resources/patient_info/{person_name}.csv').T
        X.columns = ap2_features
    except FileNotFoundError:
        return f"Data file for '{person_name}' not found."

    # Ensure the feature exists in the data
    if feat not in X.columns:
        return f"The feature '{feat}' is not found in the patient's data."

    print("works3")
    # Make a copy of the data to modify
    X_modified = X.copy()

    X_modified[feat] = value

    # Load the Autoprognosis model
    with open('./resources/model/ap_model.pkl', 'rb') as f:
        model = dill.load(f)

    # Prepare the data for prediction
    # Assuming the model expects a single sample in a DataFrame
    # If necessary, adjust the data formatting accordingly

    # Calculate the original risk score
    score_old = model.predict_proba(X)[0].round(4)  # Assuming binary classification and [0, 1] is the probability of the positive class

    # Calculate the new risk score after modification
    score_new = model.predict_proba(X_modified)[0].round(4)

    # Calculate the difference in risk
    diff = score_new - score_old

    # Determine the direction of change
    if diff > 0:
        direction = 'increased'
    elif diff < 0:
        direction = 'decreased'
    else:
        direction = 'remained the same'

    # Prepare the result string
    result = (
        f"Original CVD risk score for '{person_name}': {score_old * 100:.2f}%.\n"
        f"After changing '{feat}' to {value}, the new CVD risk score is {score_new * 100:.2f}%.\n"
        f"The risk has {direction} by {abs(diff) * 100:.2f}%."
    )
    return result


@tool
def df_to_string(name: str) -> str:
    """Use this function for any questions about different treatment options for a patient. 
    The function takes as input an empty string and returns a string that contains the information about
    the patient. This information is information on the patient's age, sex, chest pain, and others.
    Based on this information, the language model should suggest possible treatment options specific
    to this individual."""
    df = person_data('')

    # Create a dictionary to map column names to more interpretable strings
    column_map = {
        'age': 'The age of the patient is',
        'sex': 'The sex of the patient is',
        'cp': 'The type of chest pain is',
        'trestbps': 'The resting blood pressure (in mm Hg on admission to the hospital) is',
        'chol': 'The serum cholesterol in mg/dl is',
        'fbs': 'The fasting blood sugar > 120 mg/dl is',
        'restecg': 'The resting electrocardiographic results are',
        'thalch': 'The maximum heart rate achieved is',
        'exang': 'Exercise induced angina is',
        'oldpeak': 'ST depression induced by exercise relative to rest is',
        'slope': 'The slope of the peak exercise ST segment is',
        'ca': 'The number of major vessels (0-3) colored by fluoroscopy is',
        'thal': 'Thalassemia is'
    }

    # Convert the sex and fbs columns to more interpretable strings
    df['sex'] = df['sex'].apply(lambda x: 'male' if x == 1 else 'female')
    df['fbs'] = df['fbs'].apply(lambda x: 'true' if x == 1 else 'false')

    # Convert the DataFrame to a string
    df_string = ', '.join([f'{column_map[col]} {df.iloc[0][col]}' for col in df.columns])

    return df_string

@tool
def calculate_ap_score(name: str) -> str:
    """Use this function to calculate the cardiovascular disease risk for a person / calculate the Autoprognosis score for a person. The input to the function is the person_cvd_i index, e.g. person_cvd_0, person_cvd_1, ...
    The function returns a string containing information about the Autoprognosis score of a person. Interpret the results for the user as well"""
    # Dynamic imports
    from model_utils import manual_ap2_model_convert, calib_function
    ap2_features = ['sex',  'age',  'b_atrial_fibr',  'b_steroid_treat',  'b_diab_type2',  'family_cvd',  'sbp',  'smallbin',  'alkaline_phosphatase',  'apolipoprotein_a',  'apolipoprotein_b',  'urea',  'c_reactive_protein',  'cystatin_c',  'glycated_haemoglobin',  'igf_0',  'lipoprotein_a',  'triglycerides',  'overall_health_rating',  'ht_treat']

    # Get autoprognosis model
    with open('./resources/model/ap_model.pkl', 'rb') as f:
        model = dill.load(f)
    X = pd.read_csv(f'./resources/patient_info/{name}.csv').T
    X.columns = ap2_features
    score = model.predict_proba(X)[0].round(4)

    score_level = 'High' if score > 0.15 else 'Moderate/Low'
            
    return score, f"The Autoprognosis CVD score for this person is {score * 100} % by running the full AP model. This is {score_level} risk and should be interpreted as such. Interpret the results for the user and their significance."

@tool
def calculate_Qrisk_score(name: str) -> str:
    """Use this function to calculate the cardiovascular disease risk for a person / calculate the Qrisk score for a person. The input to the function is an empty string.
    The function returns a string containing information about the Q-risk score of a person."""
    
    qrisk_model = QRisk3Model()
    X = pd.read_csv('./resources/patient_info/person_cvd.csv')  
    score = qrisk_model.predict(X).values[0][0].round(3)
            
    return f"The Qrisk Risk Score for this person is {score * 100} % by running the full Qrisk3 model. QRisk3 is the recommended CVD risk score in the UK. All the uploaded variables were used in the prediction."

@tool
def get_nice_guidelines(question: str) -> str: 
    """Use this function to get the guidelines from NICE on how to treat a person with cardiovascular disease.
    a document titled Cardiovascular disease: risk assessment and reduction, including lipid modification. This guideline provides a systematic approach for assessing cardiovascular disease (CVD) risk and offers recommendations for risk reduction, covering when to conduct formal risk assessments, how to use tools like QRISK3, and strategies for lifestyle changes, lipid management, and statin use to lower CVD risk in primary care settings.
     
    The input to the function is a question, such as "What are the guidelines for a person with a x% probability of cardiovascular disease?"
    The function returns a string containing the guidelines.
    """
    # Load and split the PDF
    doc = Document('resources/documents/NICE_guidelines.docx')
    full_text = ', '.join([p.text for p in doc.paragraphs])

    # Summarize the information
    template = """Using the following information below, you are asked to provide answers to the question asked: {question}:

    {full_text}

    Answer:"""

     # Create a prompt for the LLM
    prompt = PromptTemplate(template=template, input_variables=['question', "full_text"])
    
    # Create an LLM chain and run the summarization
    llm = AzureChatOpenAI(
        openai_api_base = config['api_base'],
        openai_api_version = config['api_version'],
        deployment_name = config['deployment_id'],
        openai_api_key = config['api_key'],
        openai_api_type = config['api_type'],
        temperature = 0)
    
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    answer = llm_chain.run({'question': question, 'full_text': full_text})
    
    return answer


@tool
def get_infosheet(question: str) -> str:
    """This QRISK3 infosheet provides a comprehensive overview of the cardiovascular risk prediction model, including its methodology, a breakdown of traditional and additional risk factors, detailed validation results, and performance metrics.
    
    The input to the function is a question you would like to ask.
    The function returns a string containing the reasons explaining the answer as well as the page number.
    """
    # Load and split the PDF
    doc = Document('resources/documents/QRISK_Infosheet.docx')

    full_text = question + ' /n' + ', '.join([p.text for p in doc.paragraphs])

    # Create a summarization template
    template = """Using the following information below, you are asked to provide answers to the question: {question}. 

    {full_text}

    In case the answer does not exist in the infosheet, state so and provide your best answer without the infosheet."""

    # Create a prompt for the LLM
    prompt = PromptTemplate(template=template, input_variables=["question", "full_text"])
    
    # Create an LLM chain and run the summarization
    llm = AzureChatOpenAI(
        openai_api_base = config['api_base'],
        openai_api_version = config['api_version'],
        deployment_name = config['deployment_id'],
        openai_api_key = config['api_key'],
        openai_api_type = config['api_type'],
        temperature = 0)
    
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    answer = llm_chain.run({'question': question, 'full_text': full_text})
    
    return answer

@tool
def get_qrisk3_information(question: str) -> str:
    """Use this tool to get information about the QRISK3 method for cardiovascular risk prediction.
    
    The input to the function is a question you would like to ask in the form of a string.

    The function returns a string containing the reasons explaining the answer as well as the page number.
    """
    # Load and split the PDF
    doc = Document('./resources/documents/QRISK_Infosheet.docx')
    full_text = ', '.join([p.text for p in doc.paragraphs])

    # Create a summarization template
    template =f"""Using the following information below, you are asked to provide answers to the question: {question}

    {full_text}

    Answer:"""

    # Create a prompt for the LLM
    prompt = PromptTemplate(template=template, input_variables=["question", "full_text"])
    
    # Create an LLM chain and run the summarization
    llm = AzureChatOpenAI(
        openai_api_base = config['api_base'],
        openai_api_version = config['api_version'],
        deployment_name = config['deployment_id'],
        openai_api_key = config['api_key'],
        openai_api_type = config['api_type'],
        temperature = 0)
    
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    answer = llm_chain.run(full_text)
    
    return answer


@tool
def get_stroke_infosheet(question: str) -> str:
    """This CHA2DS2-VASc infosheet provides a comprehensive overview of the prediction model, including its methodology, a breakdown of traditional and additional risk factors, detailed validation results, and performance metrics.
    
    The input to the function is a question you would like to ask.
    The function returns a string containing the reasons explaining the answer as well as the page number.
    """
    # Load and split the PDF
    doc = Document('resources/documents/CHA2DS2_VASc_Infosheet.docx')

    full_text = question + ' /n' + ', '.join([p.text for p in doc.paragraphs])

    # Create a summarization template
    template = """Using the following information below, you are asked to provide answers to the question: {question}. 

    {full_text}

    In case the answer does not exist in the infosheet, state so and provide your best answer without the infosheet."""

    # Create a prompt for the LLM
    prompt = PromptTemplate(template=template, input_variables=["question", "full_text"])
    
    # Create an LLM chain and run the summarization
    llm = AzureChatOpenAI(
        openai_api_base = config['api_base'],
        openai_api_version = config['api_version'],
        deployment_name = config['deployment_id'],
        openai_api_key = config['api_key'],
        openai_api_type = config['api_type'],
        temperature = 0)
    
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    answer = llm_chain.run({'question': question, 'full_text': full_text})
    
    return answer


@tool
def get_nice_guidelines_stroke(question: str) -> str: 
    """Use this function to get the guidelines from NICE on how to treat a person with atrial fibrillation
    a document titled Atrial fibrillation: diagnosis and management Overview This guideline covers diagnosing and managing atrial fibrillation in adults. It includes guidance on providing the best care and treatment for people with atrial fibrillation, including assessing and managing risks of stroke and bleeding.
     
    The input to the function is a question, such as "What are the guidelines for a person with a x% probability of Atrial fibrillation?"
    The function returns a string containing the guidelines.
    """
    # Load and split the PDF
    doc = Document('resources/documents/NICE_guidelines_stroke.docx')
    full_text = ', '.join([p.text for p in doc.paragraphs])

    # Summarize the information
    template = """Using the following information below, you are asked to provide answers to the question asked: {question}:

    {full_text}

    Answer:"""

     # Create a prompt for the LLM
    prompt = PromptTemplate(template=template, input_variables=['question', "full_text"])
    
    # Create an LLM chain and run the summarization
    llm = AzureChatOpenAI(
        openai_api_base = config['api_base'],
        openai_api_version = config['api_version'],
        deployment_name = config['deployment_id'],
        openai_api_key = config['api_key'],
        openai_api_type = config['api_type'],
        temperature = 0)
    
    llm_chain = LLMChain(prompt=prompt, llm=llm)
    answer = llm_chain.run({'question': question, 'full_text': full_text})
    
    return answer

@tool
def get_feature_importance_af(name: str) -> str:
    """
    Use this for questions when asking for factors of the risk FOR A SPECIFIC person.

    The input to this function is the index of the person (e.g. 'person_af_0', 'person_af_0', etc). 
    The function returns the top factors contributing to their risk.

    """

    # Load the input data for the specified person
    # Assuming the person's data is saved as 'name.csv'
    X = pd.read_csv(f"./resources/patient_info/{name.lower()}.csv").T.reset_index()
    features = ["congestive_heart_failure", "hypertension", "age",
                "diabetes", "stroke", "vascular_disease",
                "sex"
                ]

    X.columns = features

    feature_importances = [0] * len(features)
    if X["congestive_heart_failure"].iloc[0]:
        feature_importances[0] = 1

    if X["hypertension"].iloc[0]:
        feature_importances[1] = 1

    if X["age"].iloc[0]>=75:
        feature_importances[2] = 2
    elif 65<=X["age"].iloc[0]<75:
        feature_importances[2] = 1

    if X["diabetes"].iloc[0]:
        feature_importances[3] = 1

    if X["stroke"].iloc[0]:
        feature_importances[4] = 2

    if X["vascular_disease"].iloc[0]:
        feature_importances[5] = 1

    if X["sex"].iloc[0]==0: # Female
        feature_importances[6] = 1

    # Filter and sort
    filtered = [(v, i) for v, i in zip(feature_importances, features) if v > 0]
    sorted_features = [i for v, i in sorted(filtered, key=lambda x: x[0], reverse=True)]

    return str(sorted_features)

def inference(
    congestive_heart_failure: bool,
    hypertension: bool,
    age: float,
    diabetes: bool,
    stroke: bool,
    vascular_disease: bool,
    sex: int,
) -> int:
    """Return raw CHA₂DS₂-VASc score (0-9)."""
    score = 0
    if congestive_heart_failure: score += 1
    if hypertension:             score += 1
    if age >= 75:                score += 2
    elif 65 <= age < 75:         score += 1
    if diabetes:                 score += 1
    if stroke:                   score += 2
    if vascular_disease:         score += 1
    if sex == 0:                 score += 1      # female
    return score


_SCORE_TO_PERCENT = {
    0: 0.2, 1: 0.6, 2: 2.2, 3: 3.2, 4: 4.8,
    5: 7.2, 6: 9.7, 7: 11.2, 8: 11.2, 9: 12.2,
}

_EXPECTED_COLS = [
    "congestive_heart_failure", "hypertension", "age",
    "diabetes", "stroke", "vascular_disease", "sex",
]


class CHA2DS2_VASc_RiskModel:
    """Lightweight wrapper so the API mirrors a scikit-style predict()."""
    def fit(self, *args, **kwargs):
        return self

    def predict(self, df: pd.DataFrame) -> pd.DataFrame:
        # ensure required columns exist, fill missing with zeros
        df = df.copy()
        for col in _EXPECTED_COLS:
            if col not in df.columns:
                df[col] = 0
        scores = df.apply(lambda row: inference(**row[_EXPECTED_COLS]), axis=1)
        return pd.DataFrame(scores, columns=[365])   # keep your old shape


stroke_model = CHA2DS2_VASc_RiskModel()

@tool
def calculate_stroke_risk_score(name: str) -> str:
    """
    Use this function to calculate the stroke-risk score for a given patient. The  input to the function is a name (e.g. person_af_0)

    The function returns information about the patient's risk score.
    """
    # ── 1. Load the patient record ────────────────────────────────────────────
    X = pd.read_csv(f"./resources/patient_info/{name}.csv").T.reset_index()
    X.columns = _EXPECTED_COLS

    # ── 2. Run model ──────────────────────────────────────────────────────────
    score = int(stroke_model.predict(X).iloc[0, 0])
    risk_pct = _SCORE_TO_PERCENT[score]

    return f"The score is {score} and the risk percent is {risk_pct}"

@tool
def counterfactual_stroke_risk(input_data: str) -> str:
    """
    Estimate how a *single* feature change would alter a patient's annual
    ischaemic-stroke risk according to the CHA₂DS₂-VASc score. 

    **Input (one string)**
        "person_name: 'person_af_0', feature_change: ('age', 75)"

  The function changes the required characteristic to the set value and re-runs the risk prediction.
    The function takes a string in the form tuple as an input which is "(feature, value)", such as "('age', 75)". 
    The function then returns a string explaining the old and new risk predictions, as well as their difference."

    **Output (string)**
        • Original risk (%, CHA₂DS₂-VASc score)\
        • New risk after the change (%, score)\
        • Difference and direction of change
    """
    # ── 1. Parse the instruction ------------------------------------------------
    m = re.match(r"person_name:\s*'(.*?)',\s*feature_change:\s*(\(.+\))", input_data)
    if not m:
        return ("Invalid input format. Use: "
                "\"person_name: 'name', feature_change: (feature, value)\"")
    person_name = m.group(1)
    try:
        feat, new_val = ast.literal_eval(m.group(2))
    except Exception:
        return "feature_change must be a Python tuple, e.g. ('age', 70)"

    # ── 2. Load patient data -----------------------------------------------------
    try:
        X = pd.read_csv(f"./resources/patient_info/{person_name}.csv").T.reset_index()
        X.columns = _EXPECTED_COLS           # from the shared module
    except FileNotFoundError:
        return f"Data file for '{person_name}' not found."

    if feat not in X.columns:
        return f"The feature '{feat}' isn't present in this patient's record."

    # ── 3. Original stroke risk --------------------------------------------------
    score_old = int(stroke_model.predict(X).iloc[0, 0])      # raw score 0-9
    risk_old = _SCORE_TO_PERCENT[score_old]                  # annual %

    # ── 4. Counterfactual risk ---------------------------------------------------
    X_cf = X.copy()
    X_cf.at[X_cf.index[0], feat] = new_val                   # apply change
    score_new = int(stroke_model.predict(X_cf).iloc[0, 0])
    risk_new = _SCORE_TO_PERCENT[score_new]

    diff = risk_new - risk_old
    direction = ("increased" if diff > 0 else
                 "decreased" if diff < 0 else
                 "remained the same")

    # ── 5. Compose explanation ---------------------------------------------------
    return (
        f"Original annual stroke risk for '{person_name}': "
        f"**{risk_old:.1f}%** (CHA₂DS₂-VASc score {score_old}).\n"
        f"After changing **{feat}** to **{new_val}**, the new risk is "
        f"**{risk_new:.1f}%** (score {score_new}).\n"
        f"The risk has **{direction} by {abs(diff):.1f}%**."
    )